import docx

file = docx.Document()

# 添加段落
file.add_paragraph('Hello World', 'Title')  # add_paragraph 方法可以用于写入一段文段，第一个参数接受写入的内容，第二段参数接受样式。返回的是 Paragraph 对象

para_obj1 = file.add_paragraph('This is a second paragraph.')
para_obj2 = file.add_paragraph('This is a yet another paragraph.')
para_obj1.add_run("This text is being added to the second paragraph")  # para_obj1 是一个新的自然段，可以在该自然段后通过 add_run 方法，在此自然段后添加句子。 add_run 返回的是 Run 对象


# 添加标题
# add_heading 方法可以添加标题，标题层次为0-4，add_heading 接受两个参数，第一个参数是标题内容，第二个参数是标题层次。返回的是 Paragraph 对象
file.add_heading('Header 0', 0)
file.add_heading('Header 1', 1)
file.add_heading('Header 2', 2)
file.add_heading('Header 3', 3)
file.add_heading('Header 4', 4)

# 添加换页符/换行符
# file.add_break()  # 添加换行符
file.paragraphs[-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)  # 插入分页符
file.add_paragraph("This is a new page")

# 插入图片
file.paragraphs[-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
file.add_picture("gaoda.jpg")

file.save('write_doc.docx')
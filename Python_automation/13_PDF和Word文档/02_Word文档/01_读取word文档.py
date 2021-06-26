import docx

# python-docx 中有很多结构，这些结构用三种类型来表示
# 第一中为 Document： Document 对象表示整个文档，包含一个 Paragraph 对象的列表
# 第二种为 Paragraph： Paragraph 对象表示文档中的段落(即用户在 Word 文档中输入时，每按下一次回车，视为一个新段落开始)，段落中包含一个 Run 对象的列表
# 第三种为 Run： 一个 Run 对象表示同一个样式的文本的延续，当文本样式发生改变，就需要一个新的 Run 对象

# 打开 word 文档
doc = docx.Document('demo.docx')

# 查看 word 文档的 paragraphs 对象数(可以理解为 word 文档有多少行)
print(len(doc.paragraphs))
# 查看第一行文本信息
print(doc.paragraphs[0].text)
# 查看第二行文本信息
print(doc.paragraphs[1].text)

# 查看 word 文档的第一个 paragraphs 对象中的 Run 对象数量(即第一行中的样式数量)
print(len(doc.paragraphs[1].runs))
# 查看 Run 对象的文本信息
print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)
print(doc.paragraphs[1].runs[3].text)


# 一个读取整个 word 文档的函数
def get_text(filename):
    doc_file = docx.Document(filename)
    full_text = [para.text for para in doc_file.paragraphs]
    return "\n".join(full_text)


if __name__ == '__main__':
    print(get_text("demo.docx"))
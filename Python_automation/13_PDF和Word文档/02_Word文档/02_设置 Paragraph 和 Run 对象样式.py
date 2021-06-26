import docx

file = docx.Document("demo.docx")
print(len(file.paragraphs))

# 给整行设置样式
file.paragraphs[0].style = "Quote"

# 给 Run 对象设置一个样式
file.paragraphs[1].runs[0].style = "Strong"
file.paragraphs[1].runs[0].font.underline = True  # 设置 Run 对象字体格式
file.paragraphs[1].runs[2].font.double_strike = True

# Run 对象中字体属性支持的样式
# 属性 描述
# bold 文本以粗体出现
# italic 文本以斜体出现
# underline 文本带下划线
# strike 文本带删除线
# double_strike 文本带双删除线
# all_caps 文本以大写首字母出现
# small_caps 文本以大写首字母出现，小写字母小两个点
# shadow 文本带阴影
# outline 文本以轮廓线出现，而不是实心
# rtl 文本从右至左书写
# imprint 文本以刻入页面的方式出现
# emboss 文本以凸出页面的方式出现

file.save("new_demo.docx")


